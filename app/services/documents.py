"""
Document handling: extract text from an upload, split it, and retrieve the
most relevant pieces for a question.

Retrieval is a small TF-IDF style scorer written in plain Python -- no
embeddings API, no extra service, no cost. For personal-scale documents
(a few files, a few hundred pages) this works well and stays instant.
"""

import logging
import math
import re
from collections import Counter
from pathlib import Path
from typing import List, Optional, Tuple

from app.config import settings
from app.database import db

logger = logging.getLogger(__name__)


class DocumentError(Exception):
    """Raised when a file cannot be accepted or read."""


# extension -> human label
SUPPORTED = {
    ".pdf": "PDF",
    ".txt": "text",
    ".md": "markdown",
    ".csv": "CSV",
    ".json": "JSON",
    ".docx": "Word document",
    ".py": "Python file",
    ".js": "JavaScript file",
    ".java": "Java file",
    ".html": "HTML file",
    ".css": "CSS file",
    ".sql": "SQL file",
    ".log": "log file",
}

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}

STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "if", "of", "in", "on", "at", "to",
    "for", "with", "about", "from", "by", "is", "are", "was", "were", "be",
    "been", "do", "does", "did", "can", "could", "should", "would", "will",
    "how", "what", "why", "when", "where", "who", "which", "this", "that",
    "these", "those", "it", "its", "as", "my", "me", "i", "you", "your",
    "please", "tell", "explain", "give", "summarize", "summarise",
}


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------
def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("PDF support needs the 'pypdf' package.") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        raise DocumentError("That PDF could not be opened. It may be corrupted.") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise DocumentError(
                "That PDF is password protected. Please remove the password first."
            ) from exc

    parts = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(f"[page {number}]\n{text}")
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("Word support needs the 'python-docx' package.") from exc
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise DocumentError("That Word file could not be opened.") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise DocumentError("That file's text encoding could not be read.")


def extract_text(path: Path, filename: str) -> str:
    """Return the plain text of a supported file."""
    suffix = Path(filename).suffix.lower()

    if suffix in IMAGE_EXTS:
        raise DocumentError(
            "Images can't be read by this text-only model. To analyse images, "
            "switch AI_MODEL to a vision model in your .env "
            "(for example nvidia/nemotron-nano-12b-v2-vl:free)."
        )
    if suffix not in SUPPORTED:
        allowed = ", ".join(sorted(SUPPORTED))
        raise DocumentError(f"Unsupported file type '{suffix}'. Supported: {allowed}")

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        text = _extract_plain(path)

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if not text:
        raise DocumentError(
            "No text could be extracted. If this is a scanned PDF, it contains "
            "images of text rather than real text, which needs OCR."
        )
    return text


# --------------------------------------------------------------------------
# Chunking
# --------------------------------------------------------------------------
def chunk_text(text: str, size: int = None, overlap: int = None) -> List[str]:
    """Split text into overlapping chunks, preferring paragraph boundaries."""
    size = size or settings.DOC_CHUNK_CHARS
    overlap = overlap if overlap is not None else settings.DOC_CHUNK_OVERLAP

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        # A single huge paragraph gets hard-split.
        if len(para) > size:
            if current:
                chunks.append(current.strip())
                current = ""
            for i in range(0, len(para), size - overlap):
                piece = para[i:i + size].strip()
                if piece:
                    chunks.append(piece)
            continue

        if len(current) + len(para) + 2 <= size:
            current += ("\n\n" if current else "") + para
        else:
            chunks.append(current.strip())
            # Carry a little context forward.
            tail = current[-overlap:] if overlap else ""
            current = (tail + "\n\n" + para).strip() if tail else para

    if current.strip():
        chunks.append(current.strip())

    return [c for c in chunks if c]


# --------------------------------------------------------------------------
# Saving an upload
# --------------------------------------------------------------------------
def safe_filename(filename: str) -> str:
    name = Path(filename or "upload").name
    name = re.sub(r"[^A-Za-z0-9._\- ]", "_", name).strip() or "upload"
    return name[:120]


def store_upload(
    raw: bytes, filename: str, media_type: str, conversation_id: Optional[int]
) -> dict:
    """Validate, save, extract, chunk, and record an uploaded file."""
    limit = settings.MAX_UPLOAD_MB * 1024 * 1024
    if len(raw) == 0:
        raise DocumentError("That file is empty.")
    if len(raw) > limit:
        raise DocumentError(
            f"That file is {len(raw) / 1024 / 1024:.1f} MB. "
            f"The limit is {settings.MAX_UPLOAD_MB} MB."
        )

    name = safe_filename(filename)
    settings.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    # Unique path so two files with the same name never collide.
    target = settings.UPLOAD_DIR / name
    counter = 1
    while target.exists():
        target = settings.UPLOAD_DIR / f"{Path(name).stem}_{counter}{Path(name).suffix}"
        counter += 1

    target.write_bytes(raw)

    try:
        text = extract_text(target, name)
    except DocumentError:
        target.unlink(missing_ok=True)   # don't keep files we can't read
        raise

    chunks = chunk_text(text)
    record = db.add_document(
        conversation_id=conversation_id,
        filename=name,
        media_type=media_type or "",
        size_bytes=len(raw),
        chunks=chunks,
        stored_path=str(target),
    )
    logger.info("Stored '%s': %s chars in %s chunks", name, len(text), len(chunks))
    return record


def remove_document(document_id: int) -> bool:
    stored_path = db.delete_document(document_id)
    if stored_path is None:
        return False
    if stored_path:
        Path(stored_path).unlink(missing_ok=True)
    return True


# --------------------------------------------------------------------------
# Retrieval
# --------------------------------------------------------------------------
def _tokenise(text: str) -> List[str]:
    words = re.findall(r"[a-z0-9]{2,}", (text or "").lower())
    return [w for w in words if w not in STOPWORDS]


def _score_chunks(question: str, chunks: List[dict]) -> List[Tuple[float, dict]]:
    """Rank chunks by TF-IDF overlap with the question."""
    q_tokens = _tokenise(question)
    if not q_tokens or not chunks:
        return []

    tokenised = [_tokenise(c["content"]) for c in chunks]
    total = len(chunks)

    # How many chunks contain each query term (for inverse document frequency).
    doc_freq = Counter()
    q_set = set(q_tokens)
    for tokens in tokenised:
        for term in q_set.intersection(tokens):
            doc_freq[term] += 1

    scored = []
    for chunk, tokens in zip(chunks, tokenised):
        if not tokens:
            continue
        counts = Counter(tokens)
        score = 0.0
        for term in q_set:
            tf = counts.get(term, 0)
            if not tf:
                continue
            idf = math.log(1 + total / (1 + doc_freq[term]))
            score += (1 + math.log(tf)) * idf
        if score > 0:
            score /= math.sqrt(len(tokens))   # normalise for chunk length
            scored.append((score, chunk))

    scored.sort(key=lambda pair: pair[0], reverse=True)
    return scored


def build_context(conversation_id: Optional[int], question: str) -> str:
    """
    Return a text block of the most relevant document excerpts, or "".

    If the question doesn't overlap any chunk (e.g. "summarise this"), fall
    back to the opening chunks of the most recent document.
    """
    chunks = db.get_chunks_for_conversation(conversation_id)
    if not chunks:
        return ""

    budget = settings.DOC_CONTEXT_CHARS
    scored = _score_chunks(question, chunks)

    if scored:
        picked = [chunk for _, chunk in scored[:settings.DOC_MAX_CHUNKS]]
    else:
        newest_doc = chunks[0]["document_id"]
        picked = [c for c in chunks if c["document_id"] == newest_doc][
            : settings.DOC_MAX_CHUNKS
        ]

    parts, used = [], 0
    for chunk in picked:
        body = chunk["content"]
        if used + len(body) > budget:
            body = body[: max(0, budget - used)]
        if not body.strip():
            break
        parts.append(f'From "{chunk["filename"]}" (part {chunk["idx"] + 1}):\n{body}')
        used += len(body)
        if used >= budget:
            break

    if not parts:
        return ""

    names = sorted({c["filename"] for c in picked})
    return (
        "## Attached document excerpts\n"
        f"The user has uploaded: {', '.join(names)}.\n"
        "Answer using these excerpts when relevant. If the answer is not in "
        "them, say so instead of guessing.\n\n" + "\n\n---\n\n".join(parts)
    )
